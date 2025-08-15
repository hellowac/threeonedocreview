import json
import uuid
from collections import defaultdict
from collections.abc import Generator, Sequence
from datetime import date, datetime, timedelta
from io import BytesIO, StringIO
from typing import Annotated, Any, Literal
from urllib.parse import quote

import docx
from docx.oxml.ns import qn
from docx.shared import Pt
from fastapi import (
    APIRouter,
    BackgroundTasks,
    Body,
    File,
    Form,
    HTTPException,
    Path,
    Query,
    UploadFile,
)
from fastapi.responses import StreamingResponse
from loguru import logger
from sqlmodel import col, desc, func, select

from app.api.const import MEDIA_TYPE_MAP
from app.api.deps import SaveTypeDep, SessionDep, UserinfoDep
from app.api.utils import (
    download_document_from_oss_v1,
    save_document_to_local,
    save_document_to_oss_v1,
)
from app.core.config import settings
from app.crud.documents import get_or_create_project

# celery 结果
from app.models.celery_result import CeleryResult

# 其他
from app.models.common import EnumPair

# 文档及内容的模型
from app.models.documents import (
    Document,
    DocumentContent,
    DocumentContentPublic,
    DocumentContentReview,
    DocumentContentReviewPublic,
    DocumentContentReviewsPublic,
    DocumentCreate,
    DocumentPublic,
    DocumentsPublic,
    DocumentUpdate,
    Project,
    ProjectPublic,
    ProjectsPublic,
    ProjectUpdate,
)
from app.models.enums import (
    EnumType,
    EnumTypeMap,
    FileCategory,
    ProjectTypeEnum,
    ReviewStatus,
    SaveType,
    SectionPriorityMap,
    SectionTitleTypeMap,
    SectionType,
)
from app.tasks.audit import audit_docx, audit_scan_pdf
from app.tasks.reviews import review_by_agent


class DocumentEnumRoute:
    router = APIRouter(prefix="/document/enum", tags=["document enum"])

    def __init__(self) -> None:
        self.router.get("/{enum_type}", response_model=list[EnumPair])(self.read_enum)

    def read_enum(
        self, enum_type: Annotated[EnumType, Path(description="枚举类型")]
    ) -> list[EnumPair]:
        """获取枚举"""

        enum_instance = EnumTypeMap[enum_type]

        members = [
            EnumPair(name=member.name, value=member.value) for member in enum_instance
        ]

        return members


class ProjectsRoute:
    router = APIRouter(prefix="/projects", tags=["projects"])

    def __init__(self) -> None:
        self.router.get("/", response_model=ProjectsPublic)(self.read_projects)
        self.router.get("/search", response_model=Sequence[str])(self.search_projects)
        self.router.get("/{proj_id}/documents", response_model=list[DocumentPublic])(
            self.read_project_documents
        )
        self.router.get(
            "/{proj_id}/reviews", response_model=list[DocumentContentPublic]
        )(self.read_project_reviews)
        self.router.post("/{proj_id}/audit", response_model=ProjectPublic)(
            self.audit_project
        )
        self.router.get("/{proj_id}/{version}/error")(self.get_proje_version_error)

    def read_projects(
        self,
        session: SessionDep,
        uinfo: UserinfoDep,
        key: Annotated[str | None, Query(description="搜索关键字")] = None,
        proj_type: Annotated[
            ProjectTypeEnum | None, Query(description="项目/工程类型")
        ] = None,
        start_date: Annotated[
            date | None, Query(description="开始日期,如: 2025-05-12")
        ] = None,
        end_date: Annotated[
            date | None, Query(description="结束日期,如: 2025-06-14")
        ] = None,
        review_status: Annotated[
            ReviewStatus | None, Query(description="审核状态")
        ] = None,
        skip: Annotated[int, Query()] = 0,
        limit: Annotated[int, Query()] = 100,
    ) -> Any:
        """检索文档."""

        # 默认前置条件, 未删除，文档类型为”三措“
        where_statement: list[Any] = [Project.is_delete == False]  # noqa: E712

        # 非超级用户则需要按用户ID过滤
        if not uinfo.is_superuser:
            where_statement.append(Project.iscuser_id == uinfo.id)

        # 关键字搜索
        if key:
            where_statement.append(col(Project.name).contains(key))

        if proj_type:
            where_statement.append(Project.type == proj_type)

        if start_date:
            where_statement.append(Project.create_at > start_date)

        if end_date:
            # <= date 转换为 < (date+1day)
            where_statement.append(Project.create_at < (end_date + timedelta(days=1)))

        if review_status:
            where_statement.append(Project.review_status == review_status)

        count_statement = (
            select(func.count()).select_from(Project).where(*where_statement)
        )
        count = session.exec(count_statement).one()
        statement = (
            select(Project)
            .where(*where_statement)
            .order_by(desc(Project.create_at))
            .offset(skip)
            .limit(limit)
        )
        projects = session.exec(statement).all()

        return ProjectsPublic(data=projects, count=count)  # type: ignore

    def search_projects(
        self,
        session: SessionDep,
        uinfo: UserinfoDep,
        key: Annotated[str | None, Query(description="搜索关键字")] = None,
    ) -> Sequence[str]:
        """数据关键字搜索项目名称"""

        statement = select(Project.name).where(col(Project.name).contains(key))
        if not uinfo.is_superuser:
            statement = statement.where(Project.iscuser_id == uinfo.id)

        statement = statement.distinct()

        names = session.exec(statement).all()

        return names

    def read_project_documents(
        self,
        session: SessionDep,
        uinfo: UserinfoDep,
        proj_id: Annotated[uuid.UUID, Path(description="项目ID")],
        version: Annotated[int | None, Query(description="版本ID/第几次提交")] = None,
    ) -> Sequence[DocumentPublic]:
        """获取某个项目的文档列表"""

        # 默认前置条件, 未删除，项目ID为指定ID
        where_statement: list[Any] = [
            Document.is_delete == False,  # noqa: E712
            Document.proj_id == proj_id,
        ]

        if not uinfo.is_superuser:
            where_statement.append(Document.iscuser_id == uinfo.id)

        if version:
            where_statement.append(Document.proj_version == version)

        statement = select(Document).where(*where_statement)
        publics = [
            DocumentPublic.model_validate(doc) for doc in session.exec(statement).all()
        ]

        return publics

    def read_project_reviews(
        self,
        session: SessionDep,
        uinfo: UserinfoDep,
        proj_id: Annotated[uuid.UUID, Path(description="项目ID")],
        version: Annotated[int, Query(description="版本ID/第几次提交")],
    ) -> Sequence[DocumentContentPublic]:
        """获取某个项目的某个版本的核查结果"""

        # 默认前置条件, 未删除，文档类型为”三措“
        statement = select(Document.id).where(
            Document.is_delete == False,  # noqa: E712
            Document.proj_id == proj_id,
            Document.proj_version == version,
            Document.file_category == FileCategory.THREESTEP,
        )

        if not uinfo.is_superuser:
            statement = statement.where(Document.iscuser_id == uinfo.id)

        doc_id = session.exec(statement).first()

        if doc_id is None:
            raise HTTPException(404, "该项目的三措文档未找到！")

        statement2 = select(DocumentContent).where(
            DocumentContent.is_delete == False,  # noqa: E712
            DocumentContent.proj_id == proj_id,
            DocumentContent.proj_version == version,
            DocumentContent.doc_id == doc_id,
        )

        if not uinfo.is_superuser:
            statement2 = statement2.where(DocumentContent.iscuser_id == uinfo.id)

        publics = [
            DocumentContentPublic.model_validate(doc)
            for doc in session.exec(statement2).all()
        ]

        # 根据SectionPriorityMap定义的优先级进行排序
        publics = sorted(publics, key=lambda x: SectionPriorityMap[x.section])

        return publics

    def audit_project(
        self,
        session: SessionDep,
        uinfo: UserinfoDep,
        proj_id: Annotated[uuid.UUID, Path(description="项目ID")],
        payload: Annotated[ProjectUpdate, Body(description="复核时传递的值")],
    ) -> ProjectPublic:
        """复核某个项目并更新状态"""

        project = session.get(Project, proj_id)

        if not project:
            raise HTTPException(404, "该项目不存在!")

        if not uinfo.is_superuser and project.iscuser_id != uinfo.id:
            raise HTTPException(404, "该项目不存在!")

        if payload.review_status not in (
            ReviewStatus.HUMAN_REVIEW_FAILD,
            ReviewStatus.HUMAN_REVIEW_PASSED,
        ):
            raise HTTPException(400, "要更新的状态不对")

        # 更新状态
        project.review_status = payload.review_status
        project.review_done_at = datetime.now()
        project.review_percent = 100  # 百分百进度

        # 是驳回理由
        if payload.audit_feedback:
            project.audit_feedback = payload.audit_feedback

        session.add(project)
        session.commit()
        session.refresh(project)

        return ProjectPublic.model_validate(project)

    def get_proje_version_error(
        self,
        session: SessionDep,
        uinfo: UserinfoDep,
        proj_id: Annotated[uuid.UUID, Path(description="项目ID")],
        version: Annotated[int, Path(description="版本")],
    ) -> str:
        """获取项目某个版本的审查错误"""

        project = session.get(Project, proj_id)

        if not project:
            raise HTTPException(404, "该项目不存在!")

        if not uinfo.is_superuser and project.iscuser_id != uinfo.id:
            raise HTTPException(404, "该项目不存在!")

        if project.review_status != ReviewStatus.AI_REVIEW_FAILURE:
            return "无错误"

        doc_statement = select(Document).where(
            Document.proj_id == proj_id, Document.proj_version == version
        )
        doc = session.exec(doc_statement).first()

        if not doc:
            return ""

        taskresult = session.exec(
            select(CeleryResult).where(
                CeleryResult.task_id == str(doc.task_id),
                CeleryResult.name == review_by_agent.name,  # type: ignore
            )
        ).first()

        if not taskresult:
            return ""

        return taskresult.traceback or "无错误"


class DocumentRoute:
    router = APIRouter(prefix="/documents", tags=["documents"])

    def __init__(self) -> None:
        self.router.get("/{id}", response_model=DocumentPublic)(self.read_document)
        self.router.post("/", response_model=DocumentsPublic)(self.create_document)
        self.router.post("/{id}/put", response_model=DocumentPublic)(
            self.update_document
        )
        self.router.post("/{id}/delete")(self.delete_document)
        self.router.get("/{id}/download")(self.download_document)

    def read_document(
        self,
        session: SessionDep,
        id: Annotated[uuid.UUID, Path()],
    ) -> Any:
        """通过 ID 获取文档。"""
        document = session.get(Document, id)

        if not document or document.is_delete:
            raise HTTPException(status_code=404, detail="未找到文档")

        public = DocumentPublic.model_validate(document)

        return public

    def create_document(
        self,
        session: SessionDep,
        uinfo: UserinfoDep,
        save_type: SaveTypeDep,
        background_tasks: BackgroundTasks,
        proj_name: Annotated[str, Form(description="工程名称")],
        proj_type: Annotated[ProjectTypeEnum, Form(description="工程类型")],
        threeone_file: Annotated[UploadFile, File(description="三措文档")],
        other_files: Annotated[
            list[UploadFile] | None, File(description="其他文档, 可以没有")
        ] = None,
        other_files_category: Annotated[
            str | None,
            Form(description="其他文档的分类字典,可以没有"),
        ] = None,
    ) -> Any:
        """上传新文档。

        # 通过Form上传三措文档及其附件和附件的分类信息

        其中:

        other_files_category 字段 应该是其他文件名称和文件分类的字典的json字符串，比如:

        ```json
        {
            "勘察单.xlsx": "其他材料",
            "xxx可行性研究报告.docx":"可研材料"
        }
        ```
        """

        # 搞到一个list中，好处理
        other_files = other_files or []
        other_files.insert(0, threeone_file)

        other_files_category_map: dict[str | None, FileCategory] = {}

        # 加载其他文件分类字典的内容
        try:
            if other_files_category is not None:
                other_files_category_map = json.loads(other_files_category)
        except Exception as e:
            logger.exception(e)
            raise HTTPException(status_code=500, detail="解析其他文件的分类字典失败.")

        # 项目处理
        project, proj_new_created = get_or_create_project(
            session, proj_name, proj_type, uinfo.id
        )

        if not proj_new_created:
            project.version += 1
            session.add(project)
            session.commit()
            session.refresh(project)

        documents = []

        for index, uploadfile in enumerate(other_files, start=1):
            file_category = other_files_category_map.get(
                uploadfile.filename, FileCategory.OTHER
            )
            if index == 1:
                file_category = FileCategory.THREESTEP

            # 提取 docx、pdf
            file_suffix = ""
            if uploadfile.filename is not None:
                file_suffix = uploadfile.filename.rsplit(".")[-1]

            # 保存到本地
            if save_type == SaveType.LOCAL:
                save_path = save_document_to_local(uploadfile)
            else: # oss
                save_path = save_document_to_oss_v1(proj_type.name, uploadfile)

            document_in = DocumentCreate(
                file_name=uploadfile.filename,  # type: ignore
                file_suffix=file_suffix,
                file_size=uploadfile.size,  # type: ignore
                file_category=file_category,
                save_type=save_type,
                save_path=save_path,
            )

            document = Document.model_validate(
                document_in,
                update={
                    "proj_id": project.id,
                    "proj_version": project.version,
                    "iscuser_id": uinfo.id,
                },
            )
            session.add(document)
            session.commit()
            session.refresh(document)

            # 审核三措文档
            if file_category == FileCategory.THREESTEP:
                # 审核docx文件
                if file_suffix == "docx":
                    audit_docx.delay( # type: ignore
                        proj_id = str(project.id),  # uuid.UUID,
                        proj_version = project.version,
                        proj_type = project.type,
                        proj_name = project.name,
                        doc_id = str(document.id),  # uuid.UUID,
                        iscuser_id = uinfo.id,
                        filepath = save_path,
                        _save_type = save_type.value,
                    )

                # 审核pdf扫描件
                else:
                    audit_scan_pdf.delay( # type: ignore
                        proj_id = str(project.id),  # uuid.UUID,
                        proj_version = project.version,
                        proj_type = project.type,
                        proj_name = project.name,
                        doc_id = str(document.id),  # uuid.UUID,
                        iscuser_id = uinfo.id,
                        filepath = save_path,
                        _save_type = save_type.value,
                    )

            documents.append(DocumentPublic.model_validate(document))

        return DocumentsPublic(data=documents, count=len(documents))

    def update_document(
        self,
        *,
        session: SessionDep,
        uinfo: UserinfoDep,
        id: Annotated[uuid.UUID, Path()],
        document_in: Annotated[DocumentUpdate, Body()],
    ) -> Any:
        """更新文档。"""

        document = session.get(Document, id)
        if not document:
            raise HTTPException(status_code=404, detail="未找到文档")

        if not uinfo.is_superuser and document.iscuser_id != uinfo.id:
            raise HTTPException(status_code=404, detail="未找到文档")

        update_dict = document_in.model_dump(exclude_unset=True)
        document.sqlmodel_update(update_dict)
        session.add(document)
        session.commit()
        session.refresh(document)
        return document

    def delete_document(
        self,
        session: SessionDep,
        uinfo: UserinfoDep,
        id: Annotated[uuid.UUID, Path()],
    ) -> str:
        """删除文档。"""
        document = session.get(Document, id)
        if not document:
            raise HTTPException(status_code=404, detail="未找到文档")

        if not uinfo.is_superuser and document.iscuser_id != uinfo.id:
            raise HTTPException(status_code=404, detail="未找到文档")

        document.is_delete = True
        session.add(document)
        session.commit()
        session.refresh(document)
        return "文档删除成功"

    def download_document(
        self,
        session: SessionDep,
        uinfo: UserinfoDep,
        id: Annotated[uuid.UUID, Path()],
    ) -> StreamingResponse:
        document = session.get(Document, id)
        if not document:
            raise HTTPException(status_code=404, detail="未找到文档")

        if not uinfo.is_superuser and document.iscuser_id != uinfo.id:
            raise HTTPException(status_code=404, detail="未找到文档")

        if document.is_delete:
            raise HTTPException(400, "该文件已被删除")

        media_type = MEDIA_TYPE_MAP.get(document.file_suffix)

        if document.save_type == SaveType.LOCAL:
            absolute_filepath = str(settings.UPLOAD_FILES_DIR / document.save_path)

            # def iter_file_content() -> Generator[bytes, Any, None]:

            #     with open(absolute_filepath, 'rb') as fr:
            #         yield from fr

            filecontent = BytesIO(b'')
            with open(absolute_filepath, 'rb') as fr:
                filecontent.write(fr.read())

            filecontent.seek(0)

        else:
            logger.info(f"从OSS存储的文件下载, object_name: {document.save_path}")
            filecontent = download_document_from_oss_v1(document.save_path)

        headers : dict[str, str] = {}
        headers["Content-Disposition"] = f"attachment; filename={document.file_name}"

        if media_type:
            headers["Content-Type"] = media_type # "application/octet-stream"

        return StreamingResponse(filecontent, media_type=media_type)


class DocumentContentRoute:
    router = APIRouter(prefix="/document/content", tags=["document content"])

    def __init__(self) -> None:
        self.router.get(
            "/{doc_id}/section/{section}", response_model=DocumentContentPublic
        )(self.read_document_content)

    def read_document_content(
        self,
        session: SessionDep,
        uinfo: UserinfoDep,
        doc_id: Annotated[uuid.UUID, Path(description="三措文档的ID（uuid）")],
        section: Annotated[SectionType, Path(description="所属节")],
    ) -> Any:
        """读取文档的指定节点的内容"""

        statement = select(DocumentContent).where(
            DocumentContent.doc_id == doc_id, DocumentContent.section == section
        )

        if not uinfo.is_superuser:
            statement = statement.where(DocumentContent.iscuser_id == uinfo.id)

        document_content = session.exec(statement).one_or_none()

        if document_content is None:
            raise HTTPException(status_code=404, detail="未找到文档内容")

        public = DocumentContentPublic.model_validate(document_content)

        return public


class DocumentContentReviewRoute:
    router = APIRouter(
        prefix="/document/content/reviews", tags=["document content reviews"]
    )

    def __init__(self) -> None:
        self.router.get(
            "/{doc_content_id}", response_model=DocumentContentReviewsPublic
        )(self.read_document_content_reviews)
        self.router.get("/{proj_id}/all", response_model=DocumentContentReviewsPublic)(
            self.read_proj_content_reviews
        )
        self.router.get("/{proj_id}/download")(
            self.download_proj_content_reviews
        )

    def read_document_content_reviews(
        self,
        session: SessionDep,
        uinfo: UserinfoDep,
        doc_content_id: Annotated[
            uuid.UUID, Path(description="三措文档某节内容的ID（uuid）")
        ],
    ) -> DocumentContentReviewsPublic:
        """读取文档的指定节点的内容的审查详细"""

        count_statement = (
            select(func.count())
            .select_from(DocumentContentReview)
            .where(DocumentContentReview.content_id == doc_content_id)
        )

        if not uinfo.is_superuser:
            count_statement = count_statement.where(
                DocumentContent.iscuser_id == uinfo.id
            )

        count = session.exec(count_statement).one()

        statement = select(DocumentContentReview).where(
            DocumentContentReview.content_id == doc_content_id
        )

        if not uinfo.is_superuser:
            statement = statement.where(DocumentContent.iscuser_id == uinfo.id)

        dc_reviews = session.exec(statement).all()

        publics = [
            DocumentContentReviewPublic.model_validate(reivew) for reivew in dc_reviews
        ]

        return DocumentContentReviewsPublic(data=publics, count=count)

    def read_proj_content_reviews(
        self,
        session: SessionDep,
        uinfo: UserinfoDep,
        proj_id: Annotated[uuid.UUID, Path(description="项目ID")],
        version: Annotated[int, Query(description="版本ID/第几次提交")],
    ) -> DocumentContentReviewsPublic:
        """获取指定项目和版本号的所有反馈详细列表"""

        count_statement = (
            select(func.count())
            .select_from(DocumentContentReview)
            .where(
                DocumentContentReview.proj_id == proj_id,
                DocumentContentReview.proj_version == version,
            )
        )

        if not uinfo.is_superuser:
            count_statement = count_statement.where(
                DocumentContentReview.iscuser_id == uinfo.id
            )

        count = session.exec(count_statement).one()

        statement = select(DocumentContentReview).where(
            DocumentContentReview.proj_id == proj_id,
            DocumentContentReview.proj_version == version,
        )

        if not uinfo.is_superuser:
            statement = statement.where(DocumentContentReview.iscuser_id == uinfo.id)

        dc_reviews = session.exec(statement).all()

        publics = [
            DocumentContentReviewPublic.model_validate(reivew) for reivew in dc_reviews
        ]

        return DocumentContentReviewsPublic(data=publics, count=count)

    def download_proj_content_reviews(
        self,
        session: SessionDep,
        uinfo: UserinfoDep,
        proj_id: Annotated[uuid.UUID, Path(description="项目ID")],
        version: Annotated[int, Query(description="版本ID/第几次提交")],
        type: Annotated[Literal['txt', 'docx'], Query(description="下载的文件格式: txt 或 docx")]
    ) -> StreamingResponse:

        statement = select(DocumentContentReview).where(
            DocumentContentReview.proj_id == proj_id,
            DocumentContentReview.proj_version == version,
        )

        if not uinfo.is_superuser:
            statement = statement.where(DocumentContentReview.iscuser_id == uinfo.id)

        project = session.get(Project, proj_id)

        if not project:
            raise HTTPException(400, "该项目不存在")

        dc_reviews = session.exec(statement).all()

        sec_reviews: dict[SectionType, list[DocumentContentReview]] = defaultdict(list)

        for review in dc_reviews:
            sec_reviews[review.section].append(review)

        title = f'【{project.name}】的第【{version}】次提交审核结果'

        headers : dict[str, str] = {}

        if type == 'txt':

            txt_content: StringIO = self.write_txt_file(title, sec_reviews)

            media_type = MEDIA_TYPE_MAP['txt']
            headers["Content-Disposition"] = f"attachment; filename*=UTF-8''{quote(title)}.txt"
            headers["Content-Type"] = media_type # "application/octet-stream"

            return StreamingResponse(txt_content, headers=headers, media_type=media_type)

        else:
            docx_content: BytesIO = self.write_docx_file(title, sec_reviews)

            media_type = MEDIA_TYPE_MAP['docx']
            headers["Content-Disposition"] = f"attachment; filename*=UTF-8''{quote(title)}.docx"
            headers["Content-Type"] = media_type # "application/octet-stream"

            return StreamingResponse(docx_content, headers=headers, media_type=media_type)



    def write_txt_file(self, title: str, sec_reviews: dict[SectionType, list[DocumentContentReview]]) -> StringIO:

        string_io = StringIO(f'{title}\n\n')

        for sec_title, section in SectionTitleTypeMap.items():
            string_io.write(f"【{sec_title}】存在的问题")
            string_io.write("\n\n")

            reviews = sec_reviews[section]

            if not reviews:
                string_io.write("该节暂不支持审查，或审查失败！")
                string_io.write("\n")
            else:

                for index, review in enumerate(reviews, start=1):

                    string_io.write(f"问题{index}:\n")
                    string_io.write(f"{'-' * 10}\n")
                    string_io.write(f"问题: {review.question}\n")
                    string_io.write(f"建议: {review.feedback}\n")
                    string_io.write(f"依据文件: {review.reference_filename or '无'}\n")
                    string_io.write(f"依据内容: {review.reference_content or '无'}\n")
                    string_io.write("\n")

            string_io.write("\n\n")

        string_io.seek(0)

        return string_io

    def write_docx_file(self, title: str, sec_reviews: dict[SectionType, list[DocumentContentReview]]) -> BytesIO:

        bytes_io = BytesIO()

        document = docx.Document()

        # 设置默认字体为中文
        style = document.styles['Normal']
        font = style.font # type: ignore

        font.name = '宋体'
        font.size = Pt(12)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        document.add_heading(title, 0)

        for sec_title, section in SectionTitleTypeMap.items():
            # document.add_heading(f"【{sec_title}】存在的问题", level=1)
            document.add_heading(sec_title, level=1)
            # document.add_paragraph("\n")

            reviews = sec_reviews[section]

            if not reviews:
                document.add_paragraph("该节暂不支持审查，或审查失败！")
                # document.add_paragraph("\n")
            else:

                for index, review in enumerate(reviews, start=1):

                    document.add_heading(f"问题{index}", level=2)
                    document.add_paragraph(f"问题: {review.question}", style='List Bullet')
                    document.add_paragraph(f"建议: {review.feedback}", style='List Bullet')
                    document.add_paragraph(f"依据文件: {review.reference_filename or '无'}", style='List Bullet')
                    document.add_paragraph(f"依据内容: {review.reference_content or '无'}", style='List Bullet')
                    # document.add_paragraph("\n")

            # document.add_paragraph("\n\n")

        document.save(bytes_io)

        bytes_io.seek(0)

        return bytes_io

document_enum_router = DocumentEnumRoute().router
projects_router = ProjectsRoute().router
documents_router = DocumentRoute().router
document_content_router = DocumentContentRoute().router
dc_review_router = DocumentContentReviewRoute().router
